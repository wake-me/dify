"""Abstract interface for document loader implementations."""
import os
import tempfile
from urllib.parse import urlparse

import requests

from core.rag.extractor.extractor_base import BaseExtractor
from core.rag.models.document import Document


class WordExtractor(BaseExtractor):
    """
    用于加载docx文件的类。

    参数:
        file_path (str): 要加载的文件的路径。
    """

    def __init__(self, file_path: str):
        """
        初始化加载器，处理文件路径，支持从网络加载。

        参数:
            file_path (str): 文件路径，可以是本地路径或Web URL。
        """
        # 设置文件路径，并展开可能的家目录符号~
        self.file_path = file_path
        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # 如果文件是Web路径，则下载到临时文件并使用该临时文件
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            # 下载失败时抛出异常
            if r.status_code != 200:
                raise ValueError(
                    "Check the url of your file; returned status code %s"
                    % r.status_code
                )

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile()
            self.temp_file.write(r.content)
            self.file_path = self.temp_file.name
        elif not os.path.isfile(self.file_path):
            # 如果文件路径无效，抛出异常
            raise ValueError("File path %s is not a valid file or url" % self.file_path)

    def __del__(self) -> None:
        # 析构函数，确保临时文件被关闭
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def extract(self) -> list[Document]:
        """
        加载指定路径的docx文件作为单页文档。

        返回:
            list[Document]: 包含文档内容和元数据的Document对象列表。
        """
        # 导入docx模块
        from docx import Document as docx_Document

        # 加载docx文件，提取每段文字
        document = docx_Document(self.file_path)
        doc_texts = [paragraph.text for paragraph in document.paragraphs]
        content = '\n'.join(doc_texts)

        # 构建并返回Document对象
        return [Document(
            page_content=content,
            metadata={"source": self.file_path},
        )]

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """
        检查URL是否有效。

        参数:
            url (str): 待检查的URL。

        返回:
            bool: 如果URL有效则为True，否则为False。
        """
        # 解析URL并检查其组件是否完整
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)